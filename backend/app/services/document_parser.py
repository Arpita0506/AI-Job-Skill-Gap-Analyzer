import pymupdf as fitz  # PyMuPDF
import docx
import io
import re

class DocumentParser:
    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        if ext == 'pdf':
            return DocumentParser._parse_pdf(file_bytes)
        elif ext in ['docx', 'doc']:
            return DocumentParser._parse_docx(file_bytes)
        elif ext == 'txt':
            return file_bytes.decode('utf-8', errors='ignore')
        else:
            # Fallback text decoder
            try:
                return file_bytes.decode('utf-8', errors='ignore')
            except Exception:
                raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> str:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_chunks = []
            for page in doc:
                text_chunks.append(page.get_text())
            full_text = "\n".join(text_chunks)
            if not full_text.strip():
                raise ValueError("PDF file appears to be empty or image-only scanned without selectable text.")
            return DocumentParser._clean_text(full_text)
        except Exception as e:
            if "empty" in str(e).lower():
                raise e
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> str:
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
            result = "\n".join(full_text)
            if not result.strip():
                raise ValueError("DOCX document is empty.")
            return DocumentParser._clean_text(result)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    @staticmethod
    def _clean_text(text: str) -> str:
        # Normalize carriage returns and non-breaking spaces
        text = text.replace('\r\n', '\n').replace('\r', '\n').replace('\xa0', ' ')
        # Collapse excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

document_parser = DocumentParser()
